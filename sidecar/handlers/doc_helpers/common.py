"""公共样式和工具函数
从原 document_design.rs 配色方案迁移，为所有文档类型提供统一的配色方案
"""

# 专业配色方案（与原 document_design.rs 保持一致）
THEME_COLORS = {
    # 标题颜色
    "heading1": "1F4E79",       # 深蓝色
    "heading2": "2E75B6",       # 中蓝色
    "heading3": "5B9BD5",       # 浅蓝色
    # 表格颜色
    "table_header_bg": "D6E4F0",    # 表头蓝色背景
    "table_alt_row_bg": "EDF2F9",   # 交替行浅蓝色背景
    "table_border": "B4C6E7",       # 蓝灰色边框
    # 强调色
    "accent": "2E75B6",         # 中蓝色强调
    # 文字颜色
    "text_primary": "333333",   # 主文字色
    "text_secondary": "666666", # 次要文字色
    # 背景
    "bg_light": "F5F7FA",       # 浅灰背景
    "bg_white": "FFFFFF",       # 白色背景
}

# 字体配置
FONT_CONFIG = {
    "east_asian": "微软雅黑",
    "latin": "Arial",
    "mono": "Consolas",
}

# PPT 配色方案
PPT_COLOR_SCHEMES = {
    "ocean": {
        "primary": "065A82",
        "secondary": "1C7293",
        "accent": "21295C",
    },
    "midnight": {
        "primary": "1E2761",
        "secondary": "CADCFC",
        "accent": "FFFFFF",
    },
    "forest": {
        "primary": "2C5F2D",
        "secondary": "97BC62",
        "accent": "F5F5F5",
    },
    "coral": {
        "primary": "F96167",
        "secondary": "F9E795",
        "accent": "2F3C7E",
    },
    "charcoal": {
        "primary": "36454F",
        "secondary": "F2F2F2",
        "accent": "212121",
    },
}


def apply_theme(doc, theme_name="default"):
    """对 python-docx Document 应用配色方案

    根据 theme_name 从 THEME_COLORS 取色，应用到文档的标题样式、正文样式和表格样式。
    对于非 python-docx Document 对象（如 openpyxl Workbook、pptx Presentation），
    直接返回原对象（这些类型有各自的专用 helper 函数）。

    Args:
        doc: 文档对象（python-docx Document 等）
        theme_name: 配色方案名称（当前仅支持 "default"）

    Returns:
        应用了配色方案的文档对象

    Note:
        仅修改样式定义，不改变文档内容。
        样式应用是幂等的：多次调用 apply_theme 效果与一次调用相同。
    """
    # 检测文档类型：python-docx Document 有 styles 属性
    if not hasattr(doc, "styles"):
        # 非 python-docx Document，直接返回（openpyxl/pptx 有各自的样式应用方式）
        return doc

    try:
        from docx.shared import RGBColor, Pt
    except ImportError:
        # python-docx 未安装，无法应用样式
        return doc

    # 颜色辅助函数：将 "RRGGBB" 字符串转为 RGBColor
    def _hex_to_rgb(hex_str):
        return RGBColor.from_string(hex_str)

    # 根据 theme_name 选择配色（当前仅 default）
    if theme_name != "default":
        # 未知主题名，回退到 default
        theme_name = "default"

    colors = THEME_COLORS

    try:
        # 1. 修改标题样式（Heading 1/2/3）的字体颜色和大小
        style_mapping = {
            "Heading 1": (colors["heading1"], Pt(16), True),
            "Heading 2": (colors["heading2"], Pt(14), True),
            "Heading 3": (colors["heading3"], Pt(12), True),
        }

        for style_name, (color_hex, size, bold) in style_mapping.items():
            try:
                style = doc.styles[style_name]
                style.font.color.rgb = _hex_to_rgb(color_hex)
                style.font.size = size
                style.font.bold = bold
                # 设置东亚字体和拉丁字体
                style.font.name = FONT_CONFIG["latin"]
            except (KeyError, AttributeError):
                # 样式不存在或属性不可写，跳过
                pass

        # 2. 修改正文样式（Normal）的字体颜色和大小
        try:
            normal_style = doc.styles["Normal"]
            normal_style.font.color.rgb = _hex_to_rgb(colors["text_primary"])
            normal_style.font.size = Pt(11)
            normal_style.font.name = FONT_CONFIG["latin"]
        except (KeyError, AttributeError):
            pass

        # 3. 对已有表格应用表头背景色
        # python-docx 表格样式应用较复杂，这里仅对表格的表头单元格设置背景色
        for table in doc.tables:
            if len(table.rows) == 0:
                continue
            # 表头行（第一行）设置背景色
            for cell in table.rows[0].cells:
                try:
                    # 设置单元格底纹（通过 XML 属性）
                    from docx.oxml.ns import qn
                    from docx.oxml import OxmlElement
                    tc_pr = cell._tc.get_or_add_tcPr()
                    # 移除已有的 shd 元素
                    for existing_shd in tc_pr.findall(qn('w:shd')):
                        tc_pr.remove(existing_shd)
                    # 添加新的 shd 元素
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), colors["table_header_bg"])
                    tc_pr.append(shd)
                except (AttributeError, Exception):
                    # XML 操作失败时跳过该单元格
                    pass
    except Exception:
        # 样式应用失败不影响文档可用性，返回原对象
        pass

    return doc


def add_styled_table(doc, headers, rows, style_name="Table Grid"):
    """向文档添加专业样式的表格

    Args:
        doc: python-docx Document 对象
        headers: 表头列表
        rows: 数据行列表（二维列表）
        style_name: 表格样式名称

    Returns:
        创建的表格对象
    """
    try:
        from docx.shared import RGBColor, Pt
    except ImportError:
        # python-docx 未安装时，创建无样式的表格
        table = doc.add_table(rows=1 + len(rows), cols=len(headers), style=style_name)
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = str(header)
        for row_idx, row_data in enumerate(rows):
            row_cells = table.rows[row_idx + 1].cells
            for col_idx, cell_data in enumerate(row_data):
                row_cells[col_idx].text = str(cell_data) if cell_data is not None else ""
        return table

    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style=style_name)

    # 添加表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = str(header)
        # 表头样式：蓝色背景 + 粗体
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(THEME_COLORS["heading1"])
                run.font.size = Pt(11)

    # 添加数据行
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = str(cell_data) if cell_data is not None else ""
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    return table
