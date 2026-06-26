"""Word 文档处理器
基于 python-docx 实现 Word 文档的读取、转换、分析
精简版：仅支持 read/convert/analyze 操作
"""

import os
import html
import logging

from docx import Document


class WordHandler:
    """Word (.docx) 文档处理器（精简版：仅支持 read/convert/analyze）"""

    logger = logging.getLogger(__name__)

    # ------------------------------------------------------------------ #
    #  读取
    # ------------------------------------------------------------------ #

    def read(self, params: dict) -> dict:
        """读取 Word 文档内容

        params:
            path: 文件路径
            include_formatting: 是否包含格式信息（等价于 include_runs=true）
            include_runs: 是否提取 Run 级字符属性（字体名/字号/粗体/斜体/下划线/颜色）
            include_tables_detailed: 是否提取表格详细结构（合并单元格/列宽/行高/表格样式）
            include_sections: 是否提取节信息（页面尺寸/方向/边距）
            include_headers_footers: 是否提取页眉页脚内容
        """
        path = params.get("path", "")
        if not path:
            self.logger.error("read: 缺少文件路径")
            return {"error": "缺少文件路径"}
        if not os.path.exists(path):
            raise FileNotFoundError(path)

        self.logger.info("read: 开始读取 Word 文档, path=%s", path)

        # 解析参数：include_formatting=true 等价于 include_runs=true（向后兼容）
        include_formatting = params.get("include_formatting", False)
        include_runs = params.get("include_runs", False) or include_formatting
        include_tables_detailed = params.get("include_tables_detailed", False)
        include_sections = params.get("include_sections", False)
        include_headers_footers = params.get("include_headers_footers", False)

        doc = Document(path)

        # ------------------------------------------------------------------ #
        #  段落读取（含可选的 Run 级字符属性）
        # ------------------------------------------------------------------ #
        paragraphs = []
        for para in doc.paragraphs:
            para_info = {
                "text": para.text,
                "style": para.style.name if para.style else None,
            }

            # 提取段落级格式属性
            if include_runs:
                # 段落对齐方式
                alignment = para.alignment
                para_info["alignment"] = str(alignment).split(".")[-1] if alignment else None

                # 段落格式属性
                pf = para.paragraph_format
                if pf:
                    # 行距（可为浮点数或 None）
                    line_spacing = pf.line_spacing
                    para_info["line_spacing"] = float(line_spacing) if line_spacing is not None else None
                    # 首行缩进（EMU 转 cm）
                    para_info["first_line_indent_cm"] = self._emu_to_cm(pf.first_line_indent)
                    # 段前段后间距（Pt 转 pt）
                    para_info["space_before_pt"] = self._pt_to_float(pf.space_before)
                    para_info["space_after_pt"] = self._pt_to_float(pf.space_after)

                # Run 级字符属性
                runs_info = []
                for run in para.runs:
                    run_info = {"text": run.text}
                    font = run.font
                    if font:
                        run_info["font_name"] = font.name
                        run_info["font_size_pt"] = self._pt_to_float(font.size)
                        run_info["bold"] = font.bold
                        run_info["italic"] = font.italic
                        run_info["underline"] = font.underline
                        # 颜色：RGBColor 对象转十六进制字符串
                        try:
                            if font.color and font.color.rgb:
                                run_info["color_rgb"] = str(font.color.rgb)
                            else:
                                run_info["color_rgb"] = None
                        except Exception:
                            run_info["color_rgb"] = None
                    runs_info.append(run_info)
                para_info["runs"] = runs_info

            paragraphs.append(para_info)

        # ------------------------------------------------------------------ #
        #  表格读取（含可选的详细结构）
        # ------------------------------------------------------------------ #
        tables = []
        for table in doc.tables:
            if include_tables_detailed:
                # 详细表格结构：合并单元格、列宽、行高、表格样式
                table_data = []
                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for col_idx, cell in enumerate(row.cells):
                        cell_info = {"text": cell.text}

                        # 检测合并单元格（通过底层 XML 的 gridSpan/vMerge）
                        tc = cell._tc
                        grid_span = tc.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridSpan")
                        if grid_span is not None:
                            val = grid_span.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
                            cell_info["col_span"] = int(val) if val else 1
                        else:
                            cell_info["col_span"] = 1

                        v_merge = tc.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vMerge")
                        if v_merge is not None:
                            val = v_merge.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
                            # val="restart" 表示合并起始单元格，val 为空或 "continue" 表示被合并
                            cell_info["v_merge"] = val if val else "continue"
                        else:
                            cell_info["v_merge"] = None

                        row_data.append(cell_info)
                    table_data.append(row_data)

                # 列宽（EMU 转 cm）
                column_widths_cm = []
                for col in table.columns:
                    column_widths_cm.append(self._emu_to_cm(col.width))

                # 行高（EMU 转 cm）
                row_heights_cm = []
                for row in table.rows:
                    row_heights_cm.append(self._emu_to_cm(row.height))

                # 表格样式名
                table_style_name = None
                if table.style:
                    table_style_name = table.style.name

                tables.append({
                    "data": table_data,
                    "column_widths_cm": column_widths_cm,
                    "row_heights_cm": row_heights_cm,
                    "style": table_style_name,
                })
            else:
                # 简化表格结构（仅文本，向后兼容）
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)

        # ------------------------------------------------------------------ #
        #  节信息读取（可选）
        # ------------------------------------------------------------------ #
        sections = []
        if include_sections:
            for section in doc.sections:
                section_info = {
                    # 页面尺寸（EMU 转 cm）
                    "page_width_cm": self._emu_to_cm(section.page_width),
                    "page_height_cm": self._emu_to_cm(section.page_height),
                    # 方向（PORTRAIT/LANDSCAPE）
                    "orientation": str(section.orientation).split(".")[-1] if section.orientation else None,
                    # 边距（EMU 转 cm）
                    "left_margin_cm": self._emu_to_cm(section.left_margin),
                    "right_margin_cm": self._emu_to_cm(section.right_margin),
                    "top_margin_cm": self._emu_to_cm(section.top_margin),
                    "bottom_margin_cm": self._emu_to_cm(section.bottom_margin),
                }
                sections.append(section_info)

        # ------------------------------------------------------------------ #
        #  页眉页脚读取（可选）
        # ------------------------------------------------------------------ #
        headers_footers = []
        if include_headers_footers:
            for idx, section in enumerate(doc.sections):
                hf_info = {"section_index": idx}

                # 页眉
                header_texts = []
                if section.header and not section.header.is_linked_to_previous:
                    for para in section.header.paragraphs:
                        if para.text.strip():
                            header_texts.append(para.text)
                hf_info["header"] = "\n".join(header_texts) if header_texts else None

                # 页脚
                footer_texts = []
                if section.footer and not section.footer.is_linked_to_previous:
                    for para in section.footer.paragraphs:
                        if para.text.strip():
                            footer_texts.append(para.text)
                hf_info["footer"] = "\n".join(footer_texts) if footer_texts else None

                headers_footers.append(hf_info)

        # ------------------------------------------------------------------ #
        #  文档属性
        # ------------------------------------------------------------------ #
        props = {
            "title": doc.core_properties.title or "",
            "author": doc.core_properties.author or "",
            "created": str(doc.core_properties.created) if doc.core_properties.created else "",
            "modified": str(doc.core_properties.modified) if doc.core_properties.modified else "",
        }

        self.logger.info("read: Word 文档读取完成, path=%s, 段落数=%d, 表格数=%d", path, len(paragraphs), len(tables))

        # 构建返回结果（仅在对应参数为 true 时包含扩展字段，保持向后兼容）
        result = {
            "paragraphs": paragraphs,
            "tables": tables,
            "properties": props,
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
        }
        if include_sections:
            result["sections"] = sections
        if include_headers_footers:
            result["headers_footers"] = headers_footers
        return result

    @staticmethod
    def _emu_to_cm(emu) -> float:
        """将 EMU (English Metric Unit) 转换为 cm，None 返回 None"""
        if emu is None:
            return None
        try:
            # 1 cm = 360000 EMU
            return round(float(emu) / 360000.0, 3)
        except (TypeError, ValueError):
            return None

    @staticmethod
    def _pt_to_float(pt) -> float:
        """将 Pt 对象转换为浮点数，None 返回 None"""
        if pt is None:
            return None
        try:
            return float(pt.pt)
        except (AttributeError, TypeError, ValueError):
            return None

    # ------------------------------------------------------------------ #
    #  格式转换
    # ------------------------------------------------------------------ #

    def convert(self, params: dict) -> dict:
        """格式转换"""
        path = params.get("path", "")
        output_path = params.get("output_path", "")
        target_format = params.get("format", "md")
        if not path:
            self.logger.error("convert: 缺少源文件路径")
            return {"error": "缺少源文件路径"}
        if not os.path.exists(path):
            raise FileNotFoundError(path)

        self.logger.info("convert: 开始格式转换, path=%s, format=%s", path, target_format)

        doc = Document(path)

        if target_format in ("md", "markdown"):
            content = self._convert_to_markdown(doc)
        elif target_format == "txt":
            content = "\n".join(para.text for para in doc.paragraphs)
        elif target_format == "pdf":
            # _convert_to_pdf 返回 None（直接写文件），content 自然为 None
            content = self._convert_to_pdf(doc, output_path or os.path.splitext(path)[0] + ".pdf")
        else:
            self.logger.error("convert: 不支持的目标格式: %s", target_format)
            return {"error": f"不支持的目标格式: {target_format}"}

        if content is None:
            self.logger.info("convert: 格式转换完成, output_path=%s, format=%s", output_path, target_format)
            return {
                "path": output_path,
                "format": target_format,
                "message": f"已转换为 {target_format} 格式",
            }
        elif output_path:
            os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(content)
            self.logger.info("convert: 格式转换完成, output_path=%s, format=%s", output_path, target_format)
            return {
                "path": output_path,
                "format": target_format,
                "message": f"已转换为 {target_format} 格式",
            }
        else:
            return {
                "content": content,
                "format": target_format,
            }

    def _convert_to_markdown(self, doc: Document) -> str:
        """将 Word 文档内容转换为 Markdown"""
        lines = []
        for para in doc.paragraphs:
            style = para.style.name if para.style else ""
            text = para.text
            if not text.strip():
                continue
            if "Heading 1" in style:
                lines.append(f"# {text}")
            elif "Heading 2" in style:
                lines.append(f"## {text}")
            elif "Heading 3" in style:
                lines.append(f"### {text}")
            elif "Heading 4" in style:
                lines.append(f"#### {text}")
            elif "List" in style:
                lines.append(f"- {text}")
            else:
                lines.append(text)

        for table in doc.tables:
            lines.append("")
            for i, row in enumerate(table.rows):
                row_text = "| " + " | ".join(cell.text for cell in row.cells) + " |"
                lines.append(row_text)
                if i == 0:
                    lines.append("| " + " | ".join("---" for _ in row.cells) + " |")
            lines.append("")

        return "\n\n".join(lines)

    def _convert_to_pdf(self, doc: Document, output_path: str) -> None:
        """将 Word 文档内容转换为 PDF（使用 reportlab 渲染）"""
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

        from handlers.font_utils import register_chinese_font
        font_name = register_chinese_font()

        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)

        doc_pdf = SimpleDocTemplate(output_path, pagesize=A4)
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("CustomTitle", parent=styles["Title"], fontName=font_name, fontSize=24, spaceAfter=30)
        heading_style = ParagraphStyle("CustomHeading", parent=styles["Heading2"], fontName=font_name, fontSize=16, spaceAfter=12)
        body_style = ParagraphStyle("CustomBody", parent=styles["Normal"], fontName=font_name, fontSize=12, leading=20, spaceAfter=10)

        elements = []
        for para in doc.paragraphs:
            style = para.style.name if para.style else ""
            text = para.text
            if not text.strip():
                continue
            if "Heading 1" in style:
                elements.append(Paragraph(html.escape(text), heading_style))
            elif "Heading 2" in style:
                elements.append(Paragraph(html.escape(text), heading_style))
            elif "Heading 3" in style:
                elements.append(Paragraph(html.escape(text), body_style))
            elif "Title" in style:
                elements.append(Paragraph(html.escape(text), title_style))
            else:
                elements.append(Paragraph(html.escape(text), body_style))
            elements.append(Spacer(1, 0.3 * cm))

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                t = Table(table_data)
                t.setStyle(TableStyle([
                    ("GRID", (0, 0), (-1, -1), 0.5, "#999999"),
                    ("FONTNAME", (0, 0), (-1, -1), font_name),
                    ("FONTSIZE", (0, 0), (-1, -1), 10),
                ]))
                elements.append(t)
                elements.append(Spacer(1, 0.5 * cm))

        doc_pdf.build(elements)

    # ------------------------------------------------------------------ #
    #  分析
    # ------------------------------------------------------------------ #

    def analyze(self, params: dict) -> dict:
        """分析 Word 文档"""
        path = params.get("path", "")
        if not path:
            self.logger.error("analyze: 缺少文件路径")
            return {"error": "缺少文件路径"}
        if not os.path.exists(path):
            raise FileNotFoundError(path)

        self.logger.info("analyze: 开始分析 Word 文档, path=%s", path)

        doc = Document(path)

        total_chars = sum(len(p.text) for p in doc.paragraphs)
        total_words = sum(len(p.text.split()) for p in doc.paragraphs)
        heading_count = sum(
            1 for p in doc.paragraphs
            if p.style and ("Heading" in p.style.name or "Title" in p.style.name)
        )

        headings = []
        for para in doc.paragraphs:
            if para.style and ("Heading" in para.style.name or "Title" in para.style.name):
                level = 1
                try:
                    level = int(para.style.name.replace("Heading ", ""))
                except ValueError:
                    if "Title" in para.style.name:
                        level = 0
                headings.append({"level": level, "text": para.text})

        self.logger.info("analyze: Word 文档分析完成, path=%s, 段落数=%d, 标题数=%d", path, len(doc.paragraphs), heading_count)
        return {
            "file_size": os.path.getsize(path),
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "total_chars": total_chars,
            "total_words": total_words,
            "heading_count": heading_count,
            "headings": headings,
            "properties": {
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
                "created": str(doc.core_properties.created) if doc.core_properties.created else "",
                "modified": str(doc.core_properties.modified) if doc.core_properties.modified else "",
            },
        }
